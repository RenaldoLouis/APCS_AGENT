import docx

doc = docx.Document('APCS_System_Documentation.docx')

print("--- Paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text.strip()}")

print("\n--- Tables ---")
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}:")
    for r_idx, row in enumerate(table.rows):
        row_text = " | ".join([cell.text.replace('\n', ' ').strip() for cell in row.cells])
        print(f"Row {r_idx}: {row_text}")
