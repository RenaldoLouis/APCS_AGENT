import docx
from docx.shared import Inches
import copy
import re

def create_revised_document():
    doc = docx.Document('APCS_System_Documentation.docx')
    
    # Text replacements
    replacements = {
        "an administrative control centre and a dedicated adjudicator portal give the organising team and jury the tools to manage the event end to end.":
        "An administrative control centre allows the IT team (specifically Renaldo Louis) and the admin (hello@apcsmusic.com) to use the admin dashboard to manage the ticketing system from venue, ticket pricing, and dates when people can buy tickets so it is managed end-to-end. Additionally, a dedicated adjudicator portal gives the jury a single, dedicated platform to score the registrants.",

        "The platform is delivered as a single-page web application with a public-facing website":
        "The platform is delivered as an interactive website accessible to the public",

        "a live registration countdown, a bilingual interface (Indonesian and English), and navigation across Home, About, Gallery, Achievers, Contact, and Register.":
        "a live registration countdown, a bilingual interface (Indonesian and English), and navigation across key pages to support the business: Home (main landing page for announcements), About (details the history and mission of APCS), Gallery (showcases past event photos and videos), Achievers (highlights past winners to build credibility), Contact (provides support channels), and Register (the main portal for new participants to sign up).",

        "names and emails, YouTube performance links, age and instrument categories, sub":
        "names and emails, securely hosted performance videos (stored directly in our own AWS buckets, rather than YouTube links), age and instrument categories, sub",

        "An integrated Email Management panel, powered by Nodemailer, lets administrators send targeted communications without leaving the dashboard.":
        "An integrated Email Management panel, powered by Google Workspace (Gmail), lets administrators send targeted communications without leaving the dashboard.",

        "email is sent via Nodemailer":
        "email is sent via Google Workspace (Gmail)",

        "configure ticket tiers and dynamic pricing, create optional add-ons such as merchandise, and define sales-eligibility windows, including when public sales open.":
        "set up different ticket types and prices, add optional items for purchase like merchandise, and set the exact dates and times when people are allowed to start buying tickets.",
        
        "ASIA PACIFIC CHORAL SUMMIT": "A PIANO CONCERTO SERIES",
        "Asia Pacific Choral Summit": "A Piano Concerto Series",
        "asia pacific choral summit": "a piano concerto series"
    }

    # Apply text replacements in paragraphs
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
        
        # Point 9
        if "who use the compiled results" in p.text:
            p.text = re.sub(r'who use the compiled results.*?(?=\.|$)', "who use the compiled results to automatically generate certificates using the creative team's tools", p.text)

    # Point 6: Add paragraph
    for p in doc.paragraphs:
        if "The platform is delivered as an interactive website accessible to the public" in p.text:
            new_p = p.insert_paragraph_before("To ensure the platform remains engaging and aligned with the company’s evolving branding, the IT team updates the system annually. Because the event's theme changes every year, the IT team refreshes the design, content, and specifically the registration page information to provide a fresh experience for all registrants.")
            new_p.style = p.style

    # Point 12: Add invoice example placeholder
    for p in doc.paragraphs:
        if "7. Technical Architecture & Infrastructure Costs" in p.text:
            new_h = p.insert_paragraph_before("9. Invoice and Communication Examples")
            new_h.style = p.style # Fallback to same style to avoid KeyError, though it should be a heading
            p1 = p.insert_paragraph_before("[Placeholder: Insert example of the invoice sent via Email here]")
            p2 = p.insert_paragraph_before("[Placeholder: Insert example of the invoice sent via WhatsApp (WA) here]")
            p.insert_paragraph_before("")

    # Apply text replacements in tables (excluding the main structural changes we'll do manually)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)

    # Table 8 (Tech Stack): Update email provider
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) >= 3:
            if "Layer" in table.cell(0, 0).text and "Technology" in table.cell(0, 1).text:
                for row in table.rows:
                    if "Email" in row.cells[0].text:
                        row.cells[1].text = "Google Workspace (Gmail)"
                        row.cells[2].text = "Included in the existing Google Workspace subscription."

    # Table 0: Properly add and style the new column
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        if len(table0.columns) == 2:
            table0.add_column(Inches(2.0))
            
            table0.autofit = False
            
            # Set explicit widths so it doesn't break out of the page (Total = 6.5 inches)
            widths = [Inches(1.5), Inches(3.0), Inches(2.0)]
            for i, col in enumerate(table0.columns):
                col.width = widths[i]
            
            # Populate text and copy styles
            texts = [
                "Accessible By",
                "Public (Registrants)",
                "Public (Buyers) & Admin (Renaldo Louis, hello@apcsmusic.com)",
                "Jury Members",
                "Public (Buyers) & Admin",
                "IT Team (Renaldo Louis) & Admin (hello@apcsmusic.com)"
            ]
            
            # The original "Responsibility" column texts should also be updated based on point 4 & 5
            # We already did replace globally, but let's ensure the explicit texts from point 4 and 5 are in column 1.
            if len(table0.rows) > 1:
                table0.cell(1, 1).text = "Captures participant data, file uploads (performance videos, certificates), and entry fees. All uploaded files and data are securely saved in Google Workspace (Google Drive) ensuring no data is ever lost."
            if len(table0.rows) > 5:
                table0.cell(5, 1).text = "Central hub to manage registrants, communications, venues, deadlines, and score finalisation. This area is strictly accessible only by the IT Lead (Renaldo Louis) and the official admin account (hello@apcsmusic.com)."

            for r, row in enumerate(table0.rows):
                if r < len(texts):
                    cell_new = row.cells[2]
                    cell_old = row.cells[1]
                    
                    # Copy text
                    cell_new.text = texts[r]
                    
                    # Copy paragraph formatting (font, alignment)
                    if cell_old.paragraphs and cell_new.paragraphs:
                        p_old = cell_old.paragraphs[0]
                        p_new = cell_new.paragraphs[0]
                        p_new.style = p_old.style
                        p_new.alignment = p_old.alignment
                        # Copy run styling if any (bold, color, etc) from the first run
                        if p_old.runs and p_new.runs:
                            p_new.runs[0].bold = p_old.runs[0].bold
                            p_new.runs[0].font.color.rgb = p_old.runs[0].font.color.rgb
                            p_new.runs[0].font.size = p_old.runs[0].font.size
                            p_new.runs[0].font.name = p_old.runs[0].font.name

                    # Copy cell styling (borders, shading)
                    src_tc = cell_old._tc
                    dst_tc = cell_new._tc
                    if src_tc.tcPr is not None:
                        # Find existing tcPr or create one
                        if dst_tc.tcPr is not None:
                            dst_tc.remove(dst_tc.tcPr)
                        dst_tc.append(copy.deepcopy(src_tc.tcPr))
                
                # Apply explicit widths to every cell
                for c, cell in enumerate(row.cells):
                    cell.width = widths[c]

    doc.save('APCS_System_Documentation_Revised.docx')
    print("Fixed document saved!")

create_revised_document()
