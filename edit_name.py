import docx

def edit_document_name(input_path, output_path):
    doc = docx.Document(input_path)
    
    replacements = {
        "ASIA PACIFIC CHORAL SUMMIT": "A PIANO CONCERTO SERIES",
        "Asia Pacific Choral Summit": "A Piano Concerto Series",
        "ASIA PACIFIC CHORAL SUMMIT".lower(): "a piano concerto series"
    }

    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)

    doc.save(output_path)
    print(f"Saved further revised document to {output_path}")

edit_document_name('APCS_System_Documentation_Revised.docx', 'APCS_System_Documentation_Revised.docx')
