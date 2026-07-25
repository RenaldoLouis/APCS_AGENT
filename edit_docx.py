import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_text_in_paragraph(p, old_text, new_text):
    if old_text in p.text:
        p.text = p.text.replace(old_text, new_text)
        return True
    return False

def edit_document(input_path, output_path):
    doc = docx.Document(input_path)

    # Paragraph replacements
    replacements = {
        "an administrative control centre and a dedicated adjudicator portal give the organising team and jury the tools to manage the event end to end.":
        "An administrative control centre allows the IT team (specifically Renaldo Louis) and the admin (hello@apcsmusic.com) to use the admin dashboard to manage the ticketing system from venue, ticket pricing, and dates when people can buy tickets so it is managed end-to-end. Additionally, a dedicated adjudicator portal gives the jury a single, dedicated platform to score the registrants.",

        "The platform is delivered as a single-page web application with a public-facing website":
        "The platform is delivered as an interactive website accessible to the public",

        "a live registration countdown, a bilingual interface (Indonesian and English), and navigation across Home, About, Gallery, Achievers, Contact, and Register.":
        "a live registration countdown, a bilingual interface (Indonesian and English), and navigation across key pages to support the business: Home (main landing page for announcements), About (details the history and mission of APCS), Gallery (showcases past event photos and videos), Achievers (highlights past winners to build credibility), Contact (provides support channels), and Register (the main portal for new participants to sign up).",

        "names and emails, YouTube performance links, age and instrument categories, sub":
        "names and emails, securely hosted performance videos (stored directly in our own AWS buckets, rather than YouTube links), age and instrument categories, sub",

        "An integrated Email Management panel, powered by Nodemailer, lets administrators send targeted communications without leaving the dashboard.":
        "An integrated Email Management panel, powered by Google Workspace (Gmail), lets administrators send targeted communications without leaving the dashboard.",

        "email is sent via Nodemailer":
        "email is sent via Google Workspace (Gmail)",

        "configure ticket tiers and dynamic pricing, create optional add-ons such as merchandise, and define sales-eligibility windows, including when public sales open.":
        "set up different ticket types and prices, add optional items for purchase like merchandise, and set the exact dates and times when people are allowed to start buying tickets."
    }

    # "am, who use the compiled results as the d" might be cut off, let's just use part of it
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
        
        # Check partial for point 9
        if "who use the compiled results as the d" in p.text:
            idx = p.text.find("who use the compiled results as the d")
            # Replace rest of sentence or up to a point
            # Since we don't know the full original, let's just replace this chunk
            # wait, let me just replace "who use the compiled results as the data source" or whatever with:
            # "who use the compiled results to automatically generate certificates using the creative team's tools"
            pass

    # Actually let's search for "who use the compiled results" inside paragraphs to be safe
    for p in doc.paragraphs:
        if "who use the compiled results" in p.text:
            text = p.text
            import re
            text = re.sub(r'who use the compiled results.*?(?=\.|$)', "who use the compiled results to automatically generate certificates using the creative team's tools", text)
            p.text = text

    # Add paragraph for Point 6 after "The platform is delivered as an interactive website accessible to the public"
    # Actually, we can just append it to that paragraph or insert a new paragraph.
    for i, p in enumerate(doc.paragraphs):
        if "The platform is delivered as an interactive website accessible to the public" in p.text:
            new_p = p.insert_paragraph_before("To ensure the platform remains engaging and aligned with the company’s evolving branding, the IT team updates the system annually. Because the event's theme changes every year, the IT team refreshes the design, content, and specifically the registration page information to provide a fresh experience for all registrants.")
            new_p.style = p.style

    # Tables modification
    # Table 0: Add column and update texts
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        # Add new column header if it doesn't exist
        if len(table0.columns) == 2:
            # We can't directly append a column easily in python-docx while preserving complex layouts,
            # but we can try adding a column.
            try:
                new_col = table0.add_column(docx.shared.Inches(2.0))
                # Add Header
                table0.cell(0, 2).text = "Accessible By"
            except Exception as e:
                print(f"Failed to add column to Table 0: {e}")
            
            # Row 1 Registration Engine
            if len(table0.rows) > 1:
                table0.cell(1, 1).text = "Captures participant data, file uploads (performance videos, certificates), and entry fees. All uploaded files and data are securely saved in Google Workspace (Google Drive) ensuring no data is ever lost."
                table0.cell(1, 2).text = "Public (Registrants)"
            # Row 2 Ticketing & Seating Engine
            if len(table0.rows) > 2:
                table0.cell(2, 2).text = "Public (Buyers) & Admin (Renaldo Louis, hello@apcsmusic.com)"
            # Row 3 Scoring Engine
            if len(table0.rows) > 3:
                table0.cell(3, 2).text = "Jury Members"
            # Row 4 Payment Gateway
            if len(table0.rows) > 4:
                table0.cell(4, 2).text = "Public (Buyers) & Admin"
            # Row 5 Admin Control Centre
            if len(table0.rows) > 5:
                table0.cell(5, 1).text = "Central hub to manage registrants, communications, venues, deadlines, and score finalisation. This area is strictly accessible only by the IT Lead (Renaldo Louis) and the official admin account (hello@apcsmusic.com)."
                table0.cell(5, 2).text = "IT Team (Renaldo Louis) & Admin (hello@apcsmusic.com)"

    # Table 8 (Tech Stack): Update email provider
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) >= 3:
            if "Layer" in table.cell(0, 0).text and "Technology" in table.cell(0, 1).text:
                for row in table.rows:
                    if "Email" in row.cells[0].text:
                        row.cells[1].text = "Google Workspace (Gmail)"
                        row.cells[2].text = "Included in the existing Google Workspace subscription."

    # Point 12: Add invoice example placeholder
    # Find paragraph "Payment Processing with Paper.id" or "6. Payment Processing" and go to the end of it
    found_payment = False
    for p in doc.paragraphs:
        if "7. Technical Architecture & Infrastructure Costs" in p.text:
            new_h = p.insert_paragraph_before("9. Invoice and Communication Examples")
            new_h.style = p.style
            p1 = p.insert_paragraph_before("[Placeholder: Insert example of the invoice sent via Email here]")
            p2 = p.insert_paragraph_before("[Placeholder: Insert example of the invoice sent via WhatsApp (WA) here]")
            # add some space
            p.insert_paragraph_before("")

    doc.save(output_path)
    print(f"Saved revised document to {output_path}")

edit_document('APCS_System_Documentation.docx', 'APCS_System_Documentation_Revised.docx')
